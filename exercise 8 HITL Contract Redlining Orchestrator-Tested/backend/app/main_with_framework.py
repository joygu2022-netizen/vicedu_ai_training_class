"""
Main FastAPI Application with Multi-Agent Framework

This version uses the Agent/Team/Coordinator framework for orchestration.
Students can use this as a reference or replace main.py with this approach.
"""

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Dict, Any, List, Optional
from contextlib import asynccontextmanager
import uuid
import os
import csv
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor
# from docx.enum.text import WD_COLOR_INDEX  # Not needed, using RGBColor instead
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

# Import our multi-agent framework
from app.agents import Agent, Team, Coordinator
from app.agents.agent import ParserAgent, RiskAnalyzerAgent, RedlineGeneratorAgent
from app.agents.team import TeamPattern

# Initialize the Coordinator (manages all agents and blackboard)
coordinator = Coordinator()

# In-memory document and playbook storage (for classroom)
documents: Dict[str, Dict[str, Any]] = {}
playbooks: Dict[str, Dict[str, Any]] = {}

@asynccontextmanager
async def lifespan(app: FastAPI):
    """
    Application lifespan manager for startup and shutdown events.
    """
    # Startup
    print("Initializing multi-agent framework...")
    
    # Sequential Team: Parser -> Risk Analyzer -> Redline Generator
    sequential_team = Team(
        name="sequential_team",
        pattern=TeamPattern.SEQUENTIAL,
        description="Sequential execution: parse, assess, redline"
    )
    sequential_team.add_agent(ParserAgent())
    sequential_team.add_agent(RiskAnalyzerAgent())
    sequential_team.add_agent(RedlineGeneratorAgent())
    coordinator.register_team(sequential_team)
    
    # Manager-Worker Team
    manager_worker_team = Team(
        name="manager_worker_team",
        pattern=TeamPattern.MANAGER_WORKER,
        description="Manager decomposes work, workers execute in parallel"
    )
    # For demo, using same agents (students should implement proper manager/workers)
    manager_worker_team.add_agent(ParserAgent())
    manager_worker_team.add_agent(RiskAnalyzerAgent())
    manager_worker_team.add_agent(RedlineGeneratorAgent())
    coordinator.register_team(manager_worker_team)
    
    # Pipeline Team
    pipeline_team = Team(
        name="planner_executor_team",
        pattern=TeamPattern.PIPELINE,
        description="Pipeline execution with data passing between agents"
    )
    pipeline_team.add_agent(ParserAgent())
    pipeline_team.add_agent(RiskAnalyzerAgent())
    pipeline_team.add_agent(RedlineGeneratorAgent())
    coordinator.register_team(pipeline_team)
    
    print("Demo teams initialized successfully")
    print(f"Coordinator stats: {coordinator.get_stats()}")
    
    # Add some sample documents
    documents["doc_001"] = {
        "doc_id": "doc_001",
        "name": "Sample_NDA.md",
        "content": """# Non-Disclosure Agreement

## 1. Confidential Information
The parties agree to protect confidential information disclosed during the term of this agreement.

## 2. Obligations
Recipient shall not disclose confidential information to third parties without prior written consent.

## 3. Liability
Company shall be liable for any and all damages arising from breach of this agreement, including but not limited to direct, indirect, incidental, consequential, and punitive damages.

## 4. Term
This agreement shall remain in effect for a period of five (5) years from the date of execution."""
    }
    
    # Add sample playbook
    playbooks["playbook_001"] = {
        "playbook_id": "playbook_001",
        "name": "Standard NDA Policy",
        "rules": {
            "liability_cap": "12 months fees",
            "data_retention": "90 days post-termination",
            "indemnity_exclusions": ["force majeure", "third-party claims"]
        }
    }
    
    print("Sample documents and playbooks loaded")
    
    yield
    
    # Shutdown (if needed)
    print("Shutting down application...")

app = FastAPI(
    title="Exercise 8: HITL Contract Redlining Orchestrator (with Framework)",
    description="Multi-agent system for legal document review with HITL gates",
    version="2.0.0",
    lifespan=lifespan
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== Pydantic Models ====================

class RunRequest(BaseModel):
    doc_id: str
    agent_path: str = "sequential"  # sequential | manager_worker | planner_executor
    playbook_id: Optional[str] = None


class RiskApprovalItem(BaseModel):
    clause_id: str
    approved: bool = True
    comments: Optional[str] = None


class RiskApprovalRequest(BaseModel):
    run_id: str
    items: List[RiskApprovalItem]


class FinalApproveRequest(BaseModel):
    run_id: str
    approved_proposals: List[str]
    rejected_proposals: List[str] = []
    notes: Optional[str] = None


# ==================== Root Endpoint ====================

@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Welcome to Exercise 8: HITL Contract Redlining Orchestrator",
        "service": "exercise-8-with-framework",
        "version": "2.0.0",
        "description": "Multi-agent system for legal document review with HITL gates",
        "endpoints": {
            "health": "/health",
            "api_docs": "/docs",
            "openapi": "/openapi.json",
            "documents": "/api/documents",
            "playbooks": "/api/playbooks",
            "runs": "/api/runs",
            "teams": "/api/teams",
            "reports": "/api/reports"
        },
        "status": "running"
    }


# ==================== Health Check ====================

@app.get("/health")
async def health():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "exercise-8-with-framework",
        "coordinator_stats": coordinator.get_stats()
    }


# ==================== Document Management ====================

@app.post("/api/documents", 
          summary="Upload Document",
          description="Upload a document file for processing. Accepts text files (.txt, .md, .docx, etc.)",
          response_description="Returns document ID and filename")
async def upload_document(file: UploadFile = File(..., description="Document file to upload")):
    """
    Upload a document file for processing.
    
    **Note**: This endpoint expects a file upload, not JSON data.
    Use the 'Choose File' button in Swagger UI to select a file.
    """
    # Validate file
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    # Check file size (limit to 10MB)
    content_bytes = await file.read()
    if len(content_bytes) > 10 * 1024 * 1024:  # 10MB
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 10MB")
    
    try:
        # Reset file pointer
        await file.seek(0)
        content = content_bytes.decode("utf-8", errors="ignore")
        doc_id = f"doc_{uuid.uuid4().hex[:8]}"
        
        documents[doc_id] = {
            "doc_id": doc_id,
            "name": file.filename,
            "content": content,
            "uploaded_at": str(uuid.uuid1().time),
            "size": len(content_bytes)
        }
        
        return {
            "doc_id": doc_id, 
            "name": file.filename,
            "size": len(content_bytes),
            "message": "Document uploaded successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


@app.get("/api/documents")
async def list_documents():
    """List all documents"""
    return [
        {
            "doc_id": doc["doc_id"],
            "name": doc["name"],
            "uploaded_at": doc.get("uploaded_at")
        }
        for doc in documents.values()
    ]


# ==================== Playbook Management ====================

@app.post("/api/playbooks")
async def create_playbook(name: str, rules: Dict[str, Any]):
    """Create a playbook"""
    playbook_id = f"playbook_{uuid.uuid4().hex[:8]}"
    
    playbooks[playbook_id] = {
        "playbook_id": playbook_id,
        "name": name,
        "rules": rules
    }
    
    return {"playbook_id": playbook_id, "name": name}


@app.get("/api/playbooks")
async def list_playbooks():
    """List all playbooks"""
    return list(playbooks.values())


@app.delete("/api/playbooks/{playbook_id}")
async def delete_playbook(playbook_id: str):
    """Delete a playbook"""
    if playbook_id in playbooks:
        del playbooks[playbook_id]
        return {"status": "deleted"}
    raise HTTPException(status_code=404, detail="Playbook not found")


# ==================== Run Orchestration ====================

@app.post("/api/run")
async def start_run(request: RunRequest):
    """
    Start a new document review run using the multi-agent framework.
    
    This endpoint:
    1. Validates the document exists
    2. Gets policy rules from playbook (if specified)
    3. Starts a run using the Coordinator
    4. Returns run_id for tracking
    """
    # Validate document
    if request.doc_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc = documents[request.doc_id]
    
    # Get policy rules from playbook
    policy_rules = {}
    if request.playbook_id and request.playbook_id in playbooks:
        policy_rules = playbooks[request.playbook_id]["rules"]
    
    # Start run using coordinator
    try:
        run_id = coordinator.start_run(
            doc_id=request.doc_id,
            document_text=doc["content"],
            agent_path=request.agent_path,
            playbook_id=request.playbook_id,
            policy_rules=policy_rules
        )
        
        return {
            "run_id": run_id,
            "doc_id": request.doc_id,
            "agent_path": request.agent_path,
            "status": "running"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/run/{run_id}")
async def get_run(run_id: str):
    """
    Get run details including blackboard state.
    """
    run = coordinator.get_run(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    
    blackboard = coordinator.get_blackboard(run_id)
    
    return {
        **run,
        "history": blackboard.get("history", []),
        "assessments": blackboard.get("assessments", []),
        "proposals": blackboard.get("proposals", []),
        "score": blackboard.get("score", 0)
    }


@app.get("/api/runs")
async def list_runs():
    """List all runs"""
    return coordinator.list_runs()


# ==================== HITL Gates ====================

@app.post("/api/hitl/risk-approve")
async def risk_approve(request: RiskApprovalRequest):
    """
    Risk Gate: Human approval for high-risk clauses.
    """
    # Separate approved and rejected clauses
    approved = [item.clause_id for item in request.items if item.approved]
    rejected = [item.clause_id for item in request.items if not item.approved]
    comments = {item.clause_id: item.comments for item in request.items if item.comments}
    
    success = coordinator.approve_risk(
        run_id=request.run_id,
        approved_clauses=approved,
        rejected_clauses=rejected,
        comments=comments
    )
    
    if not success:
        raise HTTPException(status_code=404, detail="Run not found")
    
    return {"status": "approved", "run_id": request.run_id}


@app.post("/api/hitl/final-approve")
async def final_approve(request: FinalApproveRequest):
    """
    Final Gate: Human approval for all redline proposals.
    """
    success = coordinator.approve_final(
        run_id=request.run_id,
        approved_proposals=request.approved_proposals,
        rejected_proposals=request.rejected_proposals,
        notes=request.notes
    )
    
    if not success:
        raise HTTPException(status_code=404, detail="Run not found")
    
    return {"status": "approved", "run_id": request.run_id}


@app.get("/api/blackboard/{run_id}")
async def get_blackboard(run_id: str):
    """
    Get the blackboard (shared memory) for a run.
    """
    blackboard = coordinator.get_blackboard(run_id)
    if not blackboard:
        raise HTTPException(status_code=404, detail="Run not found")
    
    return blackboard


# ==================== Export ====================

@app.post("/api/export/redline")
async def export_redline(run_id: str, format: str = "docx"):
    """
    Export redlined document as DOCX with tracked changes.
    """
    # Validate format
    if format.lower() not in ["docx", "pdf", "csv"]:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported format '{format}'. Supported formats: docx, pdf, csv"
        )
    
    # Validate run_id format
    if not run_id or not isinstance(run_id, str):
        raise HTTPException(status_code=400, detail="Invalid run_id")
    
    blackboard = coordinator.get_blackboard(run_id)
    if not blackboard:
        raise HTTPException(status_code=404, detail=f"Run '{run_id}' not found")
    
    # Create exports directory if it doesn't exist
    exports_dir = "exports"
    try:
        os.makedirs(exports_dir, exist_ok=True)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create exports directory: {str(e)}")
    
    try:
        if format.lower() == "docx":
            # Generate DOCX with redlines
            docx_path = f"{exports_dir}/{run_id}_redlined.docx"
            await generate_redlined_docx(blackboard, docx_path)
            
            if not os.path.exists(docx_path):
                raise HTTPException(status_code=500, detail="Failed to generate DOCX file")
            
            return FileResponse(
                path=docx_path,
                filename=f"{run_id}_redlined.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif format.lower() == "pdf":
            # Generate PDF summary memo
            pdf_path = f"{exports_dir}/{run_id}_summary_memo.pdf"
            await generate_pdf_summary(blackboard, pdf_path)
            
            if not os.path.exists(pdf_path):
                raise HTTPException(status_code=500, detail="Failed to generate PDF file")
            
            return FileResponse(
                path=pdf_path,
                filename=f"{run_id}_summary_memo.pdf",
                media_type="application/pdf"
            )
        elif format.lower() == "csv":
            # Generate CSV decision card
            csv_path = f"{exports_dir}/{run_id}_decision_card.csv"
            await generate_csv_decision_card(blackboard, csv_path)
            
            if not os.path.exists(csv_path):
                raise HTTPException(status_code=500, detail="Failed to generate CSV file")
            
            return FileResponse(
                path=csv_path,
                filename=f"{run_id}_decision_card.csv",
                media_type="text/csv"
            )
        else:
            # This should not happen due to validation above, but just in case
            raise HTTPException(
                status_code=400, 
                detail=f"Format {format} not supported"
            )
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"Export failed: {str(e)}"
        )


async def generate_redlined_docx(blackboard: Dict[str, Any], output_path: str):
    """
    Generate a DOCX document with redlines showing proposed changes.
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Redlined Document Review', 0)
    title.alignment = 1  # Center alignment
    
    # Add metadata
    doc.add_paragraph(f"Document ID: {blackboard.get('doc_id', 'N/A')}")
    doc.add_paragraph(f"Run ID: {blackboard.get('run_id', 'N/A')}")
    doc.add_paragraph(f"Agent Path: {blackboard.get('agent_path', 'N/A')}")
    doc.add_paragraph(f"Generated: {blackboard.get('created_at', 'N/A')}")
    doc.add_paragraph("")  # Empty line
    
    # Add original document content
    doc.add_heading('Original Document', level=1)
    original_text = blackboard.get('document_text', '')
    doc.add_paragraph(original_text)
    
    # Add redlined sections
    doc.add_heading('Redlined Changes', level=1)
    
    clauses = blackboard.get('clauses', [])
    assessments = blackboard.get('assessments', [])
    proposals = blackboard.get('proposals', [])
    
    # Create assessment lookup
    assessment_lookup = {a['clause_id']: a for a in assessments}
    proposal_lookup = {p['clause_id']: p for p in proposals}
    
    for clause in clauses:
        clause_id = clause['clause_id']
        clause_text = clause['text']
        
        # Add clause header
        doc.add_heading(f"Clause: {clause_id}", level=2)
        
        # Add risk assessment
        if clause_id in assessment_lookup:
            assessment = assessment_lookup[clause_id]
            risk_level = assessment['risk_level']
            rationale = assessment['rationale']
            
            risk_para = doc.add_paragraph()
            risk_para.add_run("Risk Level: ").bold = True
            risk_run = risk_para.add_run(risk_level)
            if risk_level == "HIGH":
                risk_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            elif risk_level == "MEDIUM":
                risk_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            else:
                risk_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            
            doc.add_paragraph(f"Rationale: {rationale}")
        
        # Add original text
        doc.add_paragraph("Original Text:")
        orig_para = doc.add_paragraph(clause_text)
        orig_para.style = 'Normal'
        
        # Add proposed changes if available
        if clause_id in proposal_lookup:
            proposal = proposal_lookup[clause_id]
            doc.add_paragraph("Proposed Changes:")
            
            # Add original text with strikethrough
            orig_change_para = doc.add_paragraph()
            orig_run = orig_change_para.add_run(proposal['original_text'])
            orig_run.font.strike = True
            orig_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            # Add proposed text
            prop_para = doc.add_paragraph()
            prop_run = prop_para.add_run(proposal['proposed_text'])
            prop_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            prop_run.bold = True
            
            # Add rationale
            doc.add_paragraph(f"Rationale: {proposal['rationale']}")
            doc.add_paragraph(f"Variant: {proposal['variant']}")
        
        doc.add_paragraph("")  # Empty line between clauses
    
    # Add summary
    doc.add_heading('Review Summary', level=1)
    
    high_risk_count = sum(1 for a in assessments if a['risk_level'] == 'HIGH')
    medium_risk_count = sum(1 for a in assessments if a['risk_level'] == 'MEDIUM')
    low_risk_count = sum(1 for a in assessments if a['risk_level'] == 'LOW')
    
    doc.add_paragraph(f"Total Clauses: {len(clauses)}")
    doc.add_paragraph(f"High Risk: {high_risk_count}")
    doc.add_paragraph(f"Medium Risk: {medium_risk_count}")
    doc.add_paragraph(f"Low Risk: {low_risk_count}")
    doc.add_paragraph(f"Proposals Generated: {len(proposals)}")
    
    # Save the document
    doc.save(output_path)


async def generate_pdf_summary(blackboard: Dict[str, Any], output_path: str):
    """
    Generate a PDF summary memo with key findings and recommendations.
    """
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1,  # Center alignment
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        textColor=colors.darkblue
    )
    
    # Title
    story.append(Paragraph("Legal Document Review Summary", title_style))
    story.append(Spacer(1, 20))
    
    # Metadata
    story.append(Paragraph("Document Information", heading_style))
    metadata_data = [
        ["Document ID:", blackboard.get('doc_id', 'N/A')],
        ["Run ID:", blackboard.get('run_id', 'N/A')],
        ["Agent Path:", blackboard.get('agent_path', 'N/A')],
        ["Review Date:", blackboard.get('created_at', 'N/A')],
        ["Playbook:", blackboard.get('playbook_id', 'N/A')]
    ]
    
    metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
    metadata_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (1, 0), (1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(metadata_table)
    story.append(Spacer(1, 20))
    
    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))
    
    clauses = blackboard.get('clauses', [])
    assessments = blackboard.get('assessments', [])
    proposals = blackboard.get('proposals', [])
    
    high_risk_count = sum(1 for a in assessments if a['risk_level'] == 'HIGH')
    medium_risk_count = sum(1 for a in assessments if a['risk_level'] == 'MEDIUM')
    low_risk_count = sum(1 for a in assessments if a['risk_level'] == 'LOW')
    
    summary_text = f"""
    This document review analyzed {len(clauses)} clauses and identified {high_risk_count} high-risk, 
    {medium_risk_count} medium-risk, and {low_risk_count} low-risk clauses. 
    {len(proposals)} redline proposals were generated to address identified risks.
    """
    story.append(Paragraph(summary_text, styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Risk Assessment Summary
    story.append(Paragraph("Risk Assessment Summary", heading_style))
    
    risk_data = [["Risk Level", "Count", "Percentage"]]
    total_clauses = len(assessments)
    if total_clauses > 0:
        risk_data.extend([
            ["High Risk", str(high_risk_count), f"{(high_risk_count/total_clauses)*100:.1f}%"],
            ["Medium Risk", str(medium_risk_count), f"{(medium_risk_count/total_clauses)*100:.1f}%"],
            ["Low Risk", str(low_risk_count), f"{(low_risk_count/total_clauses)*100:.1f}%"]
        ])
    
    risk_table = Table(risk_data, colWidths=[2*inch, 1*inch, 1*inch])
    risk_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(risk_table)
    story.append(Spacer(1, 20))
    
    # Key Findings
    story.append(Paragraph("Key Findings", heading_style))
    
    if proposals:
        story.append(Paragraph("Critical Issues Requiring Attention:", styles['Heading3']))
        for i, proposal in enumerate(proposals, 1):
            clause_id = proposal['clause_id']
            rationale = proposal['rationale']
            story.append(Paragraph(f"{i}. {clause_id}: {rationale}", styles['Normal']))
    else:
        story.append(Paragraph("No critical issues identified in this review.", styles['Normal']))
    
    story.append(Spacer(1, 20))
    
    # Recommendations
    story.append(Paragraph("Recommendations", heading_style))
    
    if high_risk_count > 0:
        story.append(Paragraph("1. Address high-risk clauses immediately", styles['Normal']))
        story.append(Paragraph("2. Review and implement proposed redlines", styles['Normal']))
        story.append(Paragraph("3. Consider additional legal review for complex clauses", styles['Normal']))
    else:
        story.append(Paragraph("Document appears to be in good standing with minimal risk exposure.", styles['Normal']))
    
    story.append(Spacer(1, 20))
    
    # Footer
    story.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                          ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, alignment=2)))
    
    # Build PDF
    doc.build(story)


async def generate_csv_decision_card(blackboard: Dict[str, Any], output_path: str):
    """
    Generate a CSV decision card with structured data for analysis.
    """
    clauses = blackboard.get('clauses', [])
    assessments = blackboard.get('assessments', [])
    proposals = blackboard.get('proposals', [])
    
    # Create assessment lookup
    assessment_lookup = {a['clause_id']: a for a in assessments}
    proposal_lookup = {p['clause_id']: p for p in proposals}
    
    with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        
        # Header
        writer.writerow([
            'Clause_ID',
            'Clause_Heading', 
            'Risk_Level',
            'Risk_Rationale',
            'Has_Proposal',
            'Proposal_Rationale',
            'Proposal_Variant',
            'Original_Text_Length',
            'Proposed_Text_Length',
            'Review_Status',
            'Decision_Required'
        ])
        
        # Data rows
        for clause in clauses:
            clause_id = clause['clause_id']
            clause_heading = clause.get('heading', 'N/A')
            clause_text = clause.get('text', '')
            
            # Get assessment data
            assessment = assessment_lookup.get(clause_id, {})
            risk_level = assessment.get('risk_level', 'UNKNOWN')
            risk_rationale = assessment.get('rationale', 'N/A')
            
            # Get proposal data
            proposal = proposal_lookup.get(clause_id, {})
            has_proposal = 'Yes' if proposal else 'No'
            proposal_rationale = proposal.get('rationale', 'N/A')
            proposal_variant = proposal.get('variant', 'N/A')
            
            original_length = len(clause_text)
            proposed_length = len(proposal.get('proposed_text', '')) if proposal else 0
            
            # Determine review status and decision requirement
            if risk_level == 'HIGH':
                review_status = 'Critical'
                decision_required = 'Yes'
            elif risk_level == 'MEDIUM':
                review_status = 'Attention'
                decision_required = 'Yes'
            else:
                review_status = 'Approved'
                decision_required = 'No'
            
            writer.writerow([
                clause_id,
                clause_heading,
                risk_level,
                risk_rationale,
                has_proposal,
                proposal_rationale,
                proposal_variant,
                original_length,
                proposed_length,
                review_status,
                decision_required
            ])
        
        # Summary row
        writer.writerow([])  # Empty row
        writer.writerow(['SUMMARY', '', '', '', '', '', '', '', '', '', ''])
        
        high_risk_count = sum(1 for a in assessments if a['risk_level'] == 'HIGH')
        medium_risk_count = sum(1 for a in assessments if a['risk_level'] == 'MEDIUM')
        low_risk_count = sum(1 for a in assessments if a['risk_level'] == 'LOW')
        
        writer.writerow(['Total_Clauses', len(clauses), '', '', '', '', '', '', '', '', ''])
        writer.writerow(['High_Risk_Count', high_risk_count, '', '', '', '', '', '', '', '', ''])
        writer.writerow(['Medium_Risk_Count', medium_risk_count, '', '', '', '', '', '', '', '', ''])
        writer.writerow(['Low_Risk_Count', low_risk_count, '', '', '', '', '', '', '', '', ''])
        writer.writerow(['Proposals_Generated', len(proposals), '', '', '', '', '', '', '', '', ''])
        writer.writerow(['Decisions_Required', high_risk_count + medium_risk_count, '', '', '', '', '', '', '', '', ''])


# ==================== Reports & Metrics ====================

@app.get("/api/reports/slos")
async def get_slos():
    """
    Get SLO metrics.
    
    TODO for students: Implement actual metrics calculation
    """
    return {
        "latency_p95_ms": 4820,
        "quality_score": 94.2,
        "cost_per_doc_usd": 3.85,
        "success_rate": 0.942
    }


# ==================== Replay & Debug ====================

@app.get("/api/replay/{run_id}")
async def get_replay_data(run_id: str):
    """
    Get replay data for debugging.
    """
    blackboard = coordinator.get_blackboard(run_id)
    if not blackboard:
        raise HTTPException(status_code=404, detail="Run not found")
    
    return {
        "run_id": run_id,
        "history": blackboard.get("history", []),
        "blackboard_snapshots": []  # TODO: Implement checkpoint snapshots
    }


# ==================== Team Management (for debugging) ====================

@app.get("/api/teams")
async def list_teams():
    """List all registered teams"""
    return [
        team.get_info()
        for team in coordinator.teams.values()
    ]


@app.get("/api/teams/{team_name}")
async def get_team(team_name: str):
    """Get team details"""
    team = coordinator.get_team(team_name)
    if not team:
        raise HTTPException(status_code=404, detail="Team not found")
    
    return team.get_info()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8004)
